"""JSON 规格 → docx。用法：python make_docx.py <spec.json> <out.docx>"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main():
    spec_path, out_path = sys.argv[1], sys.argv[2]
    spec = json.loads(Path(spec_path).read_text())
    doc = Document()
    doc.add_heading(spec.get("title", "报告"), level=0)
    if spec.get("subtitle"):
        p = doc.add_paragraph(spec["subtitle"])
        p.runs[0].font.size = Pt(12)
    for sec in spec.get("sections", []):
        doc.add_heading(sec.get("heading", ""), level=1)
        for para in sec.get("paragraphs") or []:
            doc.add_paragraph(str(para))
        for b in sec.get("bullets") or []:
            doc.add_paragraph(str(b), style="List Bullet")
        t = sec.get("table")
        if t and t.get("headers"):
            headers, rows = t["headers"], t.get("rows") or []
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for c, h in enumerate(headers):
                table.rows[0].cells[c].text = str(h)
            for r, row in enumerate(rows, start=1):
                for c, val in enumerate(row[: len(headers)]):
                    table.rows[r].cells[c].text = str(val)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK {out_path}")


if __name__ == "__main__":
    main()
